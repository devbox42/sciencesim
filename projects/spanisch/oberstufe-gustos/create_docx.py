#!/usr/bin/env python3
"""Erstellt das Arbeitsblatt 'Me gusta' als Word-Dokument."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Farben
SPANISCH_ROT = RGBColor(0xC4, 0x1E, 0x3A)
GRAU = RGBColor(0x88, 0x88, 0x88)
FEHLER_ROT = RGBColor(0xC6, 0x28, 0x28)
SUCCESS_GRUEN = RGBColor(0x2A, 0x7A, 0x4B)

def set_cell_shading(cell, color):
    """Setzt Hintergrundfarbe einer Zelle."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(doc):
    """Fügt eine horizontale Linie hinzu."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)

def add_section_title(doc, text, color=SPANISCH_ROT):
    """Fügt einen Abschnittstitel hinzu."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)

def add_box(doc, title, content_func, border_color="C41E3A", bg_color="F0F0F0", title_bg=None):
    """Fügt eine Box mit Rahmen hinzu."""
    # Tabelle als Box-Container
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)

    # Rahmen setzen
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:color'), border_color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Hintergrund
    set_cell_shading(cell, bg_color)

    # Titel
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    if title_bg:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Inhalt
    content_func(cell)

    doc.add_paragraph()  # Abstand

def es(run, text):
    """Formatiert Text als Spanisch (rot, fett)."""
    run.text = text
    run.bold = True
    run.font.color.rgb = SPANISCH_ROT

doc = Document()

# Seitenränder und 2 Spalten
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Zweispaltig setzen
    sectPr = section._sectPr
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '400')  # Abstand zwischen Spalten
    sectPr.append(cols)

# Standardschrift
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(9)

# === KOPFZEILE ===
header_table = doc.add_table(rows=1, cols=2)
header_table.columns[0].width = Cm(10)
header_table.columns[1].width = Cm(8)
cell_left = header_table.cell(0, 0)
cell_right = header_table.cell(0, 1)

p = cell_left.paragraphs[0]
run = p.add_run("Spanisch Oberstufe – Me gusta")
run.font.size = Pt(9)

p = cell_right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Nombre: ____________________")
run.font.size = Pt(9)

# === TITEL ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Arbeitsblatt: Me gusta – Hobbys und Vorlieben")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run("Begleitend zum Lernpfad – in den Hefter übernehmen")
run.font.size = Pt(9)
run.font.color.rgb = GRAU

add_horizontal_line(doc)

# === KASTEN 1: Gustar = Gefallen ===
add_section_title(doc, "Kasten 1: Gustar = Gefallen")

# Grammatik-Box
p = doc.add_paragraph()
run = p.add_run("Grammatik")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = SPANISCH_ROT

p = doc.add_paragraph()
run = p.add_run("Wichtig: ")
run.bold = True
run = p.add_run("Das Verb ")
run = p.add_run("gustar")
run.bold = True
run.font.color.rgb = SPANISCH_ROT
run = p.add_run(' funktioniert wie "gefallen", NICHT wie "mögen"!')

p = doc.add_paragraph()
run = p.add_run("Deutsch: ")
run.bold = True
run = p.add_run("Mir gefällt die Musik.")

p = doc.add_paragraph()
run = p.add_run("Español: ")
run.bold = True
run = p.add_run("Me")
run.bold = True
run.font.color.rgb = SPANISCH_ROT
run = p.add_run(" gusta ")
run = p.add_run("la música")
run.bold = True
run.font.color.rgb = SPANISCH_ROT
run = p.add_run(".")

p = doc.add_paragraph()
run = p.add_run("Struktur:")
run.bold = True
p = doc.add_paragraph("Wem? → Indirektes Objektpronomen")
p.paragraph_format.left_indent = Cm(0.5)
p = doc.add_paragraph("Was? → Subjekt (bestimmt die Verbform!)")
p.paragraph_format.left_indent = Cm(0.5)

# Typische Fehler
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Typische Fehler")
run.bold = True
run.font.size = Pt(9)
run.font.color.rgb = FEHLER_ROT

fehler = [
    ("*Yo gusto música.", "Falsche Struktur!"),
    ("*Me gusta los libros.", "Plural: gustan!"),
    ("*Me gustan bailar.", "Infinitiv: gusta!"),
]
for falsch, korrektur in fehler:
    p = doc.add_paragraph()
    run = p.add_run(falsch)
    run.font.strike = True
    run.font.color.rgb = FEHLER_ROT
    run = p.add_run(f"  → {korrektur}")

# === KASTEN 2: Los pronombres ===
add_section_title(doc, "Kasten 2: Los pronombres")

# Pronomen-Tabelle
table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
headers = ["Person", "Pronomen", "Deutsch"]
data = [
    ("(A mí)", "me", "mir"),
    ("(A ti)", "te", "dir"),
    ("(A él/ella)", "le", "ihm/ihr"),
    ("(A nosotros)", "nos", "uns"),
    ("(A vosotros)", "os", "euch"),
    ("(A ellos)", "les", "ihnen"),
]

for i, h in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

for row_idx, (person, pronomen, deutsch) in enumerate(data, 1):
    table.cell(row_idx, 0).text = person
    cell = table.cell(row_idx, 1)
    cell.text = pronomen
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = SPANISCH_ROT
    table.cell(row_idx, 2).text = deutsch

# Merksatz (Recuerda)
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Recuerda")
run.bold = True
run.font.color.rgb = SUCCESS_GRUEN

p = doc.add_paragraph()
run = p.add_run("gusta")
run.bold = True
run.font.color.rgb = SPANISCH_ROT
run = p.add_run(" + Singular / Infinitiv")

p = doc.add_paragraph()
run = p.add_run("gustan")
run.bold = True
run.font.color.rgb = SPANISCH_ROT
run = p.add_run(" + Plural")

# Ejemplos
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Ejemplos:")
run.bold = True

ejemplos = [
    ("Me gusta bailar.", "Mir gefällt tanzen."),
    ("Te gustan los libros.", "Dir gefallen die Bücher."),
    ("A María le gusta el café.", "María gefällt Kaffee."),
    ("No nos gusta cocinar.", "Uns gefällt Kochen nicht."),
]
for esp, deu in ejemplos:
    p = doc.add_paragraph()
    run = p.add_run(esp)
    run.bold = True
    run.font.color.rgb = SPANISCH_ROT
    run = p.add_run(f"  {deu}")

# A mí / A ti Erklärung
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
run = p.add_run('Wozu "A mí", "A ti"...?')
run.bold = True
run.font.size = Pt(8)

p = doc.add_paragraph()
run = p.add_run("Optional")
run.italic = True
run = p.add_run(" – aber nützlich für:")
p.paragraph_format.left_indent = Cm(0.3)

p = doc.add_paragraph("• Betonung: A mí me gusta, a ti no.")
p.paragraph_format.left_indent = Cm(0.5)
p = doc.add_paragraph("• Klarstellung bei le/les: A María le gusta...")
p.paragraph_format.left_indent = Cm(0.5)

# === VOCABULARIO ===
add_section_title(doc, "Vocabulario: Hobbys y tiempo libre")

vocab_table = doc.add_table(rows=11, cols=2)
vocab_table.style = 'Table Grid'
vocab_headers = ["Español", "Deutsch"]
vocab_data = [
    ("la música", "die Musik"),
    ("el deporte", "der Sport"),
    ("leer", "lesen"),
    ("bailar", "tanzen"),
    ("cocinar", "kochen"),
    ("viajar", "reisen"),
    ("jugar al fútbol", "Fußball spielen"),
    ("ver películas", "Filme schauen"),
    ("escuchar música", "Musik hören"),
    ("salir con amigos", "mit Freunden ausgehen"),
]

for i, h in enumerate(vocab_headers):
    cell = vocab_table.cell(0, i)
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

for row_idx, (esp, deu) in enumerate(vocab_data, 1):
    vocab_table.cell(row_idx, 0).text = esp
    vocab_table.cell(row_idx, 1).text = deu

# === ÜBUNGEN ===
add_section_title(doc, "Übung 1: Pronomen einsetzen")
p = doc.add_paragraph("Setze das richtige Pronomen ein (me, te, le, nos, os, les):")
p.paragraph_format.space_after = Pt(4)

uebung1 = [
    ("_______ gusta el chocolate.", "(mir)"),
    ("_______ gustan los videojuegos.", "(dir)"),
    ("A María _______ gusta leer.", "(ihr)"),
    ("_______ gusta viajar.", "(uns)"),
    ("A Pedro y Ana _______ gustan los gatos.", "(ihnen)"),
]
for i, (satz, hint) in enumerate(uebung1, 1):
    p = doc.add_paragraph(f"{i}. {satz} {hint}")

add_section_title(doc, "Übung 2: gusta oder gustan?")
p = doc.add_paragraph("Ergänze die richtige Verbform:")
p.paragraph_format.space_after = Pt(4)

uebung2 = [
    "Me _______ la música.",
    "Te _______ los deportes.",
    "Le _______ bailar.",
    "Nos _______ las películas.",
    "Les _______ el fútbol.",
]
for i, satz in enumerate(uebung2, 1):
    p = doc.add_paragraph(f"{i}. {satz}")

add_section_title(doc, "Übung 3: Übersetze ins Spanische")
uebung3 = [
    "Mir gefällt Musik. _______________________________",
    "Dir gefallen die Filme. _______________________________",
    "Uns gefällt das Reisen. _______________________________",
]
for i, satz in enumerate(uebung3, 1):
    p = doc.add_paragraph(f"{i}. {satz}")

add_section_title(doc, "Übung 4: Fehler erklären")
p = doc.add_paragraph("Warum sind diese Sätze falsch? Erkläre kurz:")
p.paragraph_format.space_after = Pt(4)

p = doc.add_paragraph("1. ")
run = p.add_run("*Me gustan bailar.")
run.font.strike = True
run.font.color.rgb = FEHLER_ROT
p = doc.add_paragraph("   _____________________________________________")

p = doc.add_paragraph("2. ")
run = p.add_run("*Yo gusto el chocolate.")
run.font.strike = True
run.font.color.rgb = FEHLER_ROT
p = doc.add_paragraph("   _____________________________________________")

add_section_title(doc, "Freies Schreiben")
p = doc.add_paragraph("Schreibe 3 Sätze: Was magst du? Was magst du nicht?")
p.paragraph_format.space_after = Pt(6)

for i in range(1, 4):
    p = doc.add_paragraph(f"{i}. ________________________________________________")
    p.paragraph_format.space_after = Pt(8)

# === LÖSUNGEN ===
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
run = p.add_run("Lösungen: ")
run.bold = True
run.font.size = Pt(7)
run.font.color.rgb = GRAU
run = p.add_run("Ü1: me, te, le, nos, les | Ü2: gusta, gustan, gusta, gustan, gusta | Ü3: Me gusta la música. / Te gustan las películas. / Nos gusta viajar. | Ü4: 1) Infinitiv braucht gusta (Sg.) 2) gustar = gefallen, nicht mögen")
run.font.size = Pt(7)
run.font.color.rgb = GRAU

# Speichern
doc.save('/Users/jpc/claudecode/greenhouse-lernpfade/projects/spanisch/oberstufe-gustos/arbeitsblatt-gustos.docx')
print("Word-Dokument erstellt: arbeitsblatt-gustos.docx")
